#!/usr/bin/env python3
"""
YouTube Script Agent
Reads Obsidian vault scripts, learns tone & style, generates new scripts from news URLs.

Usage:
    python script_agent.py <URL>
    python script_agent.py https://example.com/some-news-article
"""

import os
import sys
import json
import re
import argparse
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor

# Force UTF-8 output on Windows terminals
if sys.stdout is not None and sys.stdout.encoding != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

import anthropic
import httpx
from bs4 import BeautifulSoup
from dotenv import load_dotenv

load_dotenv()

_http = httpx.Client(trust_env=False, follow_redirects=True, timeout=30)

# ── Configuration ─────────────────────────────────────────────────────────────

VAULT_PATH = Path(os.getenv("VAULT_PATH", "./vault"))
DOCX_PATH = Path(os.getenv("DOCX_PATH", "./output"))
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")
MODEL = "claude-opus-4-7"
# How many recent scripts to load as style examples (most recent = most relevant)
MAX_SCRIPTS = int(os.getenv("MAX_SCRIPTS", "30"))

# ── Obsidian vault ─────────────────────────────────────────────────────────────

_url_cache: dict[str, str] = {}


def _read_script(md_file: Path) -> dict | None:
    try:
        content = md_file.read_text(encoding="utf-8")
        if not content.strip():
            return None
        return {
            "filename": md_file.stem,
            "path": str(md_file.relative_to(VAULT_PATH)),
            "content": content,
            "mtime": md_file.stat().st_mtime,
        }
    except Exception:
        return None


def load_vault_scripts() -> list[dict]:
    """Load the most recent MAX_SCRIPTS .md scripts from the vault."""
    if not VAULT_PATH.exists():
        return []
    all_files = sorted(
        VAULT_PATH.glob("**/*.md"),
        key=lambda f: f.stat().st_mtime,
        reverse=True,
    )
    scripts = []
    for f in all_files[:MAX_SCRIPTS]:
        s = _read_script(f)
        if s:
            scripts.append(s)
    return list(reversed(scripts))


def _load_entity_indexes() -> dict[str, list[str]]:
    """Read entity index files created by link_vault.py.
    Returns {entity_name_lower: [script_stem, ...]}
    """
    index_dir = VAULT_PATH.parent
    result: dict[str, list[str]] = {}
    for idx_file in index_dir.glob("*.md"):
        try:
            content = idx_file.read_text(encoding="utf-8")
            stems = re.findall(r'\[\[Сценарии/([^\]]+)\]\]', content)
            if stems:
                result[idx_file.stem.lower()] = stems
        except Exception:
            pass
    return result


def load_smart_vault_scripts(pre_fetched_text: str) -> list[dict]:
    """Load topic-relevant scripts using Obsidian entity indexes.

    Matches entities from index files against pre-fetched source text,
    loads linked scripts first, fills remainder with recent ones.
    """
    if not VAULT_PATH.exists():
        return []

    text_lower = pre_fetched_text.lower()
    seen: set[str] = set()
    relevant_stems: list[str] = []
    matched_entities: list[str] = []

    # Try index-based matching first
    indexes = _load_entity_indexes()
    for entity_lower, stems in indexes.items():
        if entity_lower in text_lower:
            matched_entities.append(entity_lower)
            for stem in stems:
                if stem not in seen:
                    relevant_stems.append(stem)
                    seen.add(stem)

    # Build path lookup for fast access
    all_files = sorted(
        VAULT_PATH.glob("**/*.md"),
        key=lambda f: f.stat().st_mtime,
        reverse=True,
    )
    stem_to_path: dict[str, Path] = {f.stem: f for f in all_files}

    scripts: list[dict] = []
    loaded_paths: set[Path] = set()

    # Load entity-matched scripts (capped at MAX_SCRIPTS)
    for stem in relevant_stems:
        if len(scripts) >= MAX_SCRIPTS:
            break
        path = stem_to_path.get(stem)
        if path and path not in loaded_paths:
            s = _read_script(path)
            if s:
                scripts.append(s)
                loaded_paths.add(path)

    # Fill remainder with most recent
    for f in all_files:
        if len(scripts) >= MAX_SCRIPTS:
            break
        if f not in loaded_paths:
            s = _read_script(f)
            if s:
                scripts.append(s)
                loaded_paths.add(f)

    if matched_entities:
        print(f"    → Совпали сущности: {', '.join(matched_entities)}")
        print(f"    → Загружено {len([s for s in scripts if s['filename'] in seen])} тематических + {len(scripts) - len([s for s in scripts if s['filename'] in seen])} свежих")
    else:
        print(f"    → Тематических совпадений не найдено, загружено {len(scripts)} последних")

    scripts.sort(key=lambda s: s["mtime"])
    return scripts


def save_to_vault(filename: str, content: str) -> str:
    """Save script as .md file to Obsidian vault."""
    try:
        VAULT_PATH.mkdir(parents=True, exist_ok=True)
        # Sanitize filename
        safe_name = "".join(c for c in filename if c not in r'\/:*?"<>|').strip()
        if not safe_name:
            safe_name = f"script-{datetime.now().strftime('%Y-%m-%d-%H%M%S')}"
        filepath = VAULT_PATH / f"{safe_name}.md"
        filepath.write_text(content, encoding="utf-8")
        return f"Сохранено: {filepath}"
    except Exception as e:
        return f"Ошибка сохранения: {e}"

def save_to_docx(filename: str, content: str) -> str:
    """Save script as .docx to DOCX_PATH with basic formatting."""
    try:
        DOCX_PATH.mkdir(parents=True, exist_ok=True)
        safe_name = "".join(c for c in filename if c not in r'\/:*?"<>|').strip()
        if not safe_name:
            safe_name = f"script-{datetime.now().strftime('%Y-%m-%d-%H%M%S')}"
        filepath = DOCX_PATH / f"{safe_name}.docx"

        doc = Document()
        # Default font
        for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
            try:
                doc.styles[style_name].font.name = "Arial"
            except Exception:
                pass

        # Strip YAML frontmatter
        text = content
        if text.startswith("---\n"):
            end = text.find("\n---\n", 4)
            if end != -1:
                text = text[end + 5:]

        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue

            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            else:
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                if stripped.startswith("(вставить"):
                    run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
                    run.bold = True
                elif stripped.startswith("(") and stripped.endswith(")"):
                    run.italic = True
                elif stripped.startswith("[") and stripped.endswith("]"):
                    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        doc.save(filepath)
        return f"Word: {filepath}"
    except Exception as e:
        return f"Ошибка Word: {e}"


# ── YouTube transcript ─────────────────────────────────────────────────────────

YOUTUBE_RE = re.compile(
    r"(?:youtube\.com/(?:watch\?.*?v=|shorts/)|youtu\.be/)[\w-]+",
    re.IGNORECASE,
)

TELEGRAM_RE = re.compile(
    r"(?:t\.me|telegram\.me)/([^/\s?]+)/(\d+)",
    re.IGNORECASE,
)


def _parse_vtt(text: str) -> str:
    """Convert WebVTT subtitles to clean deduplicated plain text."""
    lines = []
    for line in text.splitlines():
        line = line.strip()
        if not line or line.startswith("WEBVTT") or line.startswith("NOTE"):
            continue
        if re.match(r"^[\d:\.]+\s*-->\s*[\d:\.]+", line):
            continue
        if re.match(r"^\d+$", line):
            continue
        line = re.sub(r"<[^>]+>", "", line).strip()
        if line:
            lines.append(line)
    # Remove consecutive duplicates (auto-captions repeat each phrase)
    result, prev = [], None
    for ln in lines:
        if ln != prev:
            result.append(ln)
            prev = ln
    return "\n".join(result)


def fetch_youtube_transcript(url: str) -> str:
    """Extract subtitles from a YouTube video using yt-dlp."""
    try:
        import yt_dlp
    except ImportError:
        return "yt-dlp не установлен. Запусти: pip install yt-dlp"

    ydl_opts = {"quiet": True, "no_warnings": True, "skip_download": True}
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)
    except Exception as e:
        return f"Ошибка yt-dlp {url}: {e}"

    title = info.get("title", "")
    duration = info.get("duration") or 0
    header = f"[YouTube] {title} ({duration // 60}:{duration % 60:02d})\n"

    subtitles = info.get("subtitles") or {}
    auto_caps = info.get("automatic_captions") or {}

    sub_url = lang_used = None
    for lang in ("ru", "en"):
        for source in (subtitles, auto_caps):
            if lang not in source:
                continue
            entries = source[lang]
            for entry in entries:
                if entry.get("ext") == "vtt":
                    sub_url, lang_used = entry["url"], lang
                    break
            if not sub_url and entries:
                sub_url, lang_used = entries[0]["url"], lang
            if sub_url:
                break
        if sub_url:
            break

    if not sub_url:
        desc = (info.get("description") or "")[:800]
        return header + f"(субтитры недоступны)\n\nОписание:\n{desc}"

    try:
        resp = _http.get(sub_url)
        resp.raise_for_status()
        transcript = _parse_vtt(resp.text)
    except Exception as e:
        return header + f"(ошибка загрузки субтитров: {e})"

    result = header + f"Транскрипт ({lang_used}):\n\n{transcript}"
    if len(result) > 12_000:
        result = result[:12_000] + "\n\n[... транскрипт обрезан ...]"
    return result


# ── Telegram fetcher ───────────────────────────────────────────────────────────

def fetch_telegram_post(url: str) -> str:
    """Fetch text from a public Telegram channel post."""
    m = TELEGRAM_RE.search(url)
    if not m:
        return f"Не удалось разобрать Telegram-ссылку: {url}"

    channel, post_id = m.group(1), m.group(2)
    web_url = f"https://t.me/s/{channel}/{post_id}"

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0 Safari/537.36"
        )
    }

    try:
        response = _http.get(web_url, headers=headers)
        response.raise_for_status()

        soup = BeautifulSoup(response.text, "html.parser")

        # Channel title
        title_el = soup.find(class_="tgme_channel_info_header_title")
        channel_title = title_el.get_text(strip=True) if title_el else f"@{channel}"

        # Find the specific post by data-post attribute
        post_div = soup.find(attrs={"data-post": f"{channel}/{post_id}"})

        if post_div:
            text_el = post_div.find(class_="tgme_widget_message_text")
            text = text_el.get_text(separator="\n", strip=True) if text_el else ""
            date_el = post_div.find("time")
            date_str = (date_el.get("datetime", "")[:10] + " · ") if date_el else ""
        else:
            # Fallback: grab all message texts on the page
            texts = [
                el.get_text(separator="\n", strip=True)
                for el in soup.find_all(class_="tgme_widget_message_text")
            ]
            text = "\n\n---\n\n".join(texts) if texts else ""
            date_str = ""

        if not text:
            return f"[Telegram] {channel_title} — пост #{post_id} (только медиа, текста нет)"

        return f"[Telegram] {channel_title} · {date_str}пост #{post_id}\n\n{text}"

    except Exception as e:
        return f"Ошибка загрузки Telegram-поста {url}: {e}"


# ── Web fetcher ────────────────────────────────────────────────────────────────

def fetch_url(url: str) -> str:
    """Fetch readable text from a URL. YouTube → transcript via yt-dlp."""
    if url in _url_cache:
        return _url_cache[url]
    if YOUTUBE_RE.search(url):
        result = fetch_youtube_transcript(url)
        _url_cache[url] = result
        return result
    if TELEGRAM_RE.search(url):
        result = fetch_telegram_post(url)
        _url_cache[url] = result
        return result
    try:
        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/124.0 Safari/537.36"
            )
        }
        response = _http.get(url, headers=headers)
        response.raise_for_status()

        soup = BeautifulSoup(response.text, "html.parser")

        # Remove boilerplate tags
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
            tag.decompose()

        # Prefer semantic content containers
        main = (
            soup.find("article")
            or soup.find("main")
            or soup.find(id=lambda x: x and "content" in x.lower() if x else False)
            or soup.body
        )
        raw_text = (main or soup).get_text(separator="\n", strip=True)

        # Collapse blank lines
        lines = [l.strip() for l in raw_text.splitlines() if l.strip()]
        text = "\n".join(lines)

        # Cap at ~10 000 chars
        if len(text) > 10_000:
            text = text[:10_000] + "\n\n[... текст обрезан ...]"

        result = text or "Не удалось извлечь текст со страницы."
        _url_cache[url] = result
        return result
    except Exception as e:
        return f"Ошибка загрузки URL {url}: {e}"

# ── Prompt builder ─────────────────────────────────────────────────────────────

def build_system_prompt(scripts: list[dict]) -> str:
    """Build the system prompt. All existing scripts are embedded and cached."""

    examples_block = ""
    if scripts:
        examples_block = "\n\n---\n## ТВОИ ПРЕДЫДУЩИЕ СЦЕНАРИИ\n\n"
        examples_block += (
            "Используй эти сценарии в двух целях:\n\n"
            "**1. Стиль и структура** — изучи тон, голос, формат вставок, способ подачи материала. "
            "Пиши точно так же.\n\n"
            "**2. Контекст и история** — эти сценарии показывают, что ты уже рассказывал об этой теме или артисте. "
            "Используй факты из них как фоновые знания. "
            "Не повторяй события и темы, которые уже были подробно разобраны — если только новость не добавляет к ним что-то принципиально новое. "
            "Если повторяешь что-то из прошлых роликов — делай это кратко, как напоминание, а не как основной материал.\n\n"
        )
        for s in scripts:
            examples_block += f"### {s['path']}\n\n{s['content']}\n\n---\n\n"
    else:
        examples_block = (
            "\n\n---\n## ПРИМЕРЫ СЦЕНАРИЕВ\n\n"
            "Vault пуст. Используй правила формата ниже и пиши в стиле канала «Убиваю БПМ».\n\n"
        )

    return f"""Ты — AI-агент по написанию сценариев для YouTube-канала «Убиваю БПМ».

## РОЛЬ
Канал посвящён российскому и мировому хип-хопу, рэп-культуре, стримерам и музыкальной индустрии.
Ты пишешь сценарии строго в стиле и голосе автора канала, опираясь на его предыдущие работы.

## АЛГОРИТМ РАБОТЫ
При получении URL новости ты:
1. Вызываешь fetch_url, чтобы получить текст статьи.
2. Анализируешь факты, контекст, угол подачи.
3. Пишешь полный сценарий для YouTube-ролика в стиле автора.
4. Вызываешь save_to_vault, чтобы сохранить сценарий как атомарную заметку.
5. Сообщаешь, что готово, и где сохранён файл.

## СТРОГИЕ ПРАВИЛА ФАКТЧЕКИНГА (нарушение недопустимо)

- **Только факты из источников.** Используй ТОЛЬКО информацию, которую ты получил через fetch_url. Не добавляй факты, цитаты, события или детали, которых нет в загруженных материалах.
- **Нет домыслам.** Не придумывай высказывания артистов, даты, суммы, детали биографии, названия треков, конфликты — ничего, чего нет в источнике.
- **Нет «возможно» и «вероятно».** Не строй предположений, которые могут выглядеть как факты.
- **Если факта нет — скажи об этом.** Вместо выдуманной детали используй плейсхолдер: [УТОЧНИТЬ: что именно нужно проверить]. Лучше честный пробел, чем ложь.
- **Цитаты — только дословные.** Цитируй артиста только если его слова есть в источнике дословно. Перефразировать нельзя так, чтобы это выглядело как прямая цитата.

## СТРУКТУРА СЦЕНАРИЯ (строго соблюдать)

### 1. Хук (2–4 предложения)
Зацепляющее открытие темы. Без «Всем привет» — сразу к сути. Создаёт интригу.
Пример: «Когда-то Моргенштерн был человеком, который построил карьеру на одном простом принципе: делай громко, нагло, вызывающе — и тебя заметят.»

### 2. Сигнал интро канала
`(интро убиваю бпм)`

### 3. Основная часть
Анализ, факты, цитаты. Вставки расставляются прямо по тексту по правилам ниже.

### 4. Завершение
«Что ж, поехали!» / «Что ж, давайте приступать!» / «Что ж, приступим!»

---

## ПРАВИЛА ВСТАВОК (точно соблюдать формат)

**Вставить видео целиком:**
`(вставить https://ссылка )`

**Вставить конкретный отрезок:**
`(вставить https://ссылка 02:27 – 02:32)`

**Вставить с инструкцией монтажёру:**
`(вставить https://ссылка первые 13 сек, мат убрать)`
`(вставить https://ссылка 00:00-00:45, мат убрать, ускорить на 1,1х)`

**Фон (видео/картинка под закадровый текст):**
`(вставить на фоне https://ссылка )`
`(вставить на фоне https://ссылка скрин чтоб было видно [что именно])`
`(вставить на фоне https://ссылка слайдшоу из 2 скринов)`

**Несколько источников в одной вставке:**
`(вставить на фоне https://ссылка1 + https://ссылка2 )`

**Плейсхолдер для футажа (когда ссылка будет позже):**
`[ВИДЕОРЯД НОВЫЙ ПЕРЕСЛАННЫЙ В ВК]`
`(трек … перешлю в тг)`

**Важно:** Если ссылка на источник неизвестна, используй плейсхолдер в квадратных скобках:
`[ВСТАВИТЬ: описание нужного видео/скрина]`

---

## ТОН И СТИЛЬ

- **Голос:** разговорный + аналитический, живой, прямой. Не сухой и не пафосный.
- **Обращение:** «вы», «ребята», «мы с вами». Используй «мы» когда говоришь от лица канала.
- **Цитаты:** оформляются в русских кавычках «вот так»
- **Маркеры речи:** «Дело в том, что», «Надо понимать, что», «Вот в чём дело», «Конечно», «Разумеется»
- **Ссылки на индустрию:** «наша рэп-игра», «наша индустрия», «рэп-индустрия»
- **Артисты с иностранными именами:** транслитерируй на русский (21 Сэвэдж, Найн Майс, Мэйби Бэйби)
- **Иноагенты:** при упоминании Оксимирона, Фэйса — добавлять «иноагент»/«признан иноагентом»
- **Имя файла:** формат YYYY-MM-DD Название темы (без спецсимволов)

---

## СТОП-СЛОП: ЗАПРЕЩЁННЫЕ AI-ПАТТЕРНЫ

Перед финальной версией сценария убедись, что ни одного из этих паттернов нет.

### Запрещённые фразы-затравки (throat-clearing)
Не используй фразы, которые объявляют о том, что ты собираешься сказать, вместо того чтобы сказать:
- «В этом видео мы рассмотрим...», «Сегодня я расскажу...», «Давайте разберёмся...»
- «Вот в чём дело:» / «Дело вот в чём:» — как самостоятельные вводные (не как маркер речи внутри мысли)
- «Позвольте объяснить», «Поговорим о...», «Начнём с того, что»
- Пустые переходы: «Итак», «Таким образом», «Как мы видим», «Подводя итог»

### Запрещённые усилители без содержания
- Наречия: «абсолютно», «буквально», «реально», «серьёзно», «невероятно», «откровенно говоря»
- Пустые интенсификаторы: «очень», «крайне», «исключительно» — если следующее слово не несёт конкретики
- Фальшивый пафос: «это меняет всё», «история не знает таких случаев», «это революция»

### Запрещённые структурные паттерны
- **Бинарный контраст:** «Это не X. Это Y.» — скажи Y напрямую
- **Отрицательный список:** сначала перечислять чем что-то НЕ является, потом называть
- **Драматическая фрагментация:** отдельные слова. В отдельных предложениях. Для пафоса. — запрещено
- **Риторическая подводка:** «А что если я скажу вам, что...», «Подумайте об этом:», «Представьте себе...»
- **Пассивный залог без субъекта:** «было сделано», «оказалось решено» — называй кто сделал
- **Мета-комментарий:** не ссылайся на структуру самого видео («как мы увидим далее», «в конце ролика»)

### Запрещённые декларативы
Не объявляй о важности — показывай фактами:
- «Последствия оказались значительными» → назови конкретные последствия
- «Это очень важно понять» → скажи что именно и почему
- «Ситуация непростая» → опиши конкретную сложность

### Ритм
- Не пиши три предложения подряд одинаковой длины
- Чередуй: короткий удар → развёрнутая мысль → снова удар
- Каждое предложение — живой человек что-то делает, а не абстракция «происходит»

{examples_block}"""

# ── Tool definitions (JSON Schema) ────────────────────────────────────────────

TOOLS = [
    {
        "name": "fetch_url",
        "description": (
            "Загружает текстовое содержимое веб-страницы по URL. "
            "Используй для получения текста новостной статьи перед написанием сценария."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "URL страницы для загрузки"}
            },
            "required": ["url"],
        },
    },
    {
        "name": "save_to_vault",
        "description": (
            "Сохраняет готовый сценарий как .md файл в Obsidian vault. "
            "Вызывай только когда сценарий полностью написан."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Имя файла без расширения, формат: YYYY-MM-DD Название темы",
                },
                "content": {
                    "type": "string",
                    "description": "Полное содержимое сценария в markdown",
                },
            },
            "required": ["filename", "content"],
        },
    },
]

# ── Tool executor ──────────────────────────────────────────────────────────────

def run_tool(name: str, inputs: dict) -> str:
    if name == "fetch_url":
        return fetch_url(inputs["url"])
    if name == "save_to_vault":
        md_result = save_to_vault(inputs["filename"], inputs["content"])
        docx_result = save_to_docx(inputs["filename"], inputs["content"])
        return f"{md_result}\n{docx_result}"
    return f"Неизвестный инструмент: {name}"

# ── Agent loop ─────────────────────────────────────────────────────────────────

_LOG_PATH = Path(__file__).parent / "debug.log"

def _log(text: str) -> None:
    try:
        with open(_LOG_PATH, "a", encoding="utf-8") as f:
            f.write(text)
    except Exception:
        pass

def generate_script(
    urls: list[str],
    length_min: int | None = None,
    inserts: list[str] | None = None,
    note: str | None = None,
    stop_event=None,
) -> None:
    _LOG_PATH.write_text("", encoding="utf-8")  # clear log at start
    client = anthropic.Anthropic(
        api_key=ANTHROPIC_API_KEY,
        http_client=httpx.Client(trust_env=False),
    )

    print(f"\n[>] Источников: {len(urls)}")
    for u in urls:
        print(f"    {u}")
    if length_min:
        print(f"    Длина: ~{length_min} мин")
    if inserts:
        print(f"    Вставок задано: {len(inserts)}")
    if note:
        print(f"    Указание: {note}")
    print("[*] Загружаю источники и читаю vault...")
    pre_text = ""
    for u in urls:
        pre_text += fetch_url(u) + "\n\n"

    total = sum(1 for _ in VAULT_PATH.glob("**/*.md")) if VAULT_PATH.exists() else 0
    scripts = load_smart_vault_scripts(pre_text)
    if scripts:
        print(f"    → Итого загружено {len(scripts)} из {total} сценариев")
    else:
        print("    → Vault пуст. Работаю без примеров стиля.")

    system_prompt = build_system_prompt(scripts)

    # Build user message with pre-fetched content already included
    user_msg = "Напиши полный сценарий для YouTube-ролика на основе следующих материалов.\n\n"

    for i, u in enumerate(urls, 1):
        content = fetch_url(u)
        user_msg += f"## Источник {i}: {u}\n\n{content}\n\n---\n\n"

    if length_min:
        approx_words = length_min * 140
        user_msg += f"Длина сценария: примерно {length_min} минут (~{approx_words} слов).\n\n"
    if inserts:
        user_msg += "Обязательно включи эти YouTube-вставки в сценарий:\n"
        for url in inserts:
            user_msg += f"  (вставить {url} )\n"
        user_msg += "\n"
    if note:
        user_msg += (
            f"\n⚠️ ОБЯЗАТЕЛЬНОЕ ТЗ — выполнить строго и полностью:\n"
            f"{note}\n\n"
        )

    user_msg += "Напиши полноценный сценарий и сохрани его в vault инструментом save_to_vault."

    messages = [{"role": "user", "content": user_msg}]

    print("\n[AI] Агент работает...\n")
    print("─" * 60)

    all_text_parts: list[str] = []
    vault_saved = False

    # Agentic loop
    while True:
        if stop_event and stop_event.is_set():
            print("\n[СТОП] Генерация прервана пользователем.")
            break
        with client.messages.stream(
            model=MODEL,
            max_tokens=16000,
            thinking={"type": "adaptive"},
            system=[
                {
                    "type": "text",
                    "text": system_prompt,
                    # Cache the large system prompt (scripts) to avoid re-tokenising on every turn
                    "cache_control": {"type": "ephemeral"},
                }
            ],
            tools=TOOLS,
            messages=messages,
        ) as stream:
            # Stream text blocks to terminal in real time
            for event in stream:
                if (
                    hasattr(event, "type")
                    and event.type == "content_block_delta"
                    and hasattr(event.delta, "type")
                    and event.delta.type == "text_delta"
                ):
                    print(event.delta.text, end="", flush=True)
                    all_text_parts.append(event.delta.text)

            response = stream.get_final_message()

        assistant_content = response.content
        messages.append({"role": "assistant", "content": assistant_content})

        # Done
        if response.stop_reason == "end_turn":
            print("\n" + "─" * 60)
            # Debug: show what text blocks the model returned
            for blk in assistant_content:
                if hasattr(blk, "type"):
                    print(f"[DEBUG] block type={blk.type}, len={len(getattr(blk, 'text', '') or getattr(blk, 'thinking', '') or '')}")
            break

        # Execute tools
        if response.stop_reason == "tool_use":
            print()  # newline after any streamed text
            tool_results = []

            for block in assistant_content:
                if block.type != "tool_use":
                    continue

                short_input = json.dumps(block.input, ensure_ascii=False)
                if len(short_input) > 120:
                    short_input = short_input[:120] + "…"
                print(f"\n[tool] {block.name}({short_input})")

                result = run_tool(block.name, block.input)

                if block.name == "save_to_vault":
                    vault_saved = True
                    print(f"[OK] {result}")
                else:
                    preview = result[:200].replace("\n", " ")
                    print(f"    → {preview}…" if len(result) > 200 else f"    → {preview}")

                tool_results.append(
                    {
                        "type": "tool_result",
                        "tool_use_id": block.id,
                        "content": result,
                    }
                )

            print()
            messages.append({"role": "user", "content": tool_results})
        else:
            # Unexpected stop
            break

    # Debug info
    full_text = "".join(all_text_parts).strip()
    debug_msg = f"\n[DEBUG] vault_saved={vault_saved}, text_len={len(full_text)}\n[DEBUG] text_preview={full_text[:300]!r}\n"
    print(debug_msg)
    _log(debug_msg)
    _log(f"\n--- FULL CAPTURED TEXT ---\n{full_text}\n--- END ---\n")

    # Fallback: if agent never called save_to_vault but generated text — save it automatically
    if not vault_saved:
        if len(full_text) > 1500:
            # Try to extract title from first heading or first line
            title_match = re.search(r"^#+\s*(.+)", full_text, re.MULTILINE)
            if title_match:
                raw_title = title_match.group(1).strip()
            else:
                raw_title = full_text.splitlines()[0][:60].strip()
            filename = f"{datetime.now().strftime('%Y-%m-%d')} {raw_title}"
            print(f"\n[!] save_to_vault не был вызван. Автосохранение...")
            md_result = save_to_vault(filename, full_text)
            docx_result = save_to_docx(filename, full_text)
            print(f"[OK] {md_result}")
            print(f"[OK] {docx_result}")

    print("\n[OK] Готово.")

# ── Entry point ────────────────────────────────────────────────────────────────

def main():
    if not ANTHROPIC_API_KEY:
        print("Ошибка: не задан ANTHROPIC_API_KEY.")
        print("Создай файл .env в папке script_agent/ и добавь:")
        print("  ANTHROPIC_API_KEY=sk-ant-...")
        sys.exit(1)

    parser = argparse.ArgumentParser(
        description="YouTube Script Agent — генератор сценариев для канала «Убиваю БПМ»"
    )
    parser.add_argument(
        "urls", nargs="+", metavar="URL",
        help="Один или несколько URL источников (новости, YouTube-видео)",
    )
    parser.add_argument(
        "--length", type=int, metavar="MIN",
        help="Желаемая длина сценария в минутах (например: 5, 10, 15)",
    )
    parser.add_argument(
        "--insert", nargs="+", metavar="URL",
        help="YouTube-ссылки, которые обязательно вставить в сценарий",
    )
    parser.add_argument(
        "--note", metavar="TEXT",
        help="Любые дополнительные указания агенту (в кавычках)",
    )
    args = parser.parse_args()

    generate_script(args.urls, length_min=args.length, inserts=args.insert, note=args.note)


if __name__ == "__main__":
    main()
